import os
import io
import re
import shutil
import tempfile
from datetime import datetime
from zipfile import ZipFile
import zipfile

from flask import Flask, render_template_string, request, send_file, jsonify, redirect, url_for
import pandas as pd
from docx import Document
from docx.enum.text import WD_LINE_SPACING
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Side, PatternFill, Font
from openpyxl.utils import get_column_letter

app = Flask(__name__)

# =========================================================
# Configuração geral
# =========================================================
MODELOS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'modelos')

MODELO_NACIONAL = os.path.join(MODELOS_DIR, 'Modelo Carta Pleito Backbone Nacional.docx')
MODELO_REGIONAL = os.path.join(MODELOS_DIR, 'Modelo Carta Pleito Backbone Regional.docx')

# Modelo de croqui: fica salvo no servidor (pasta modelos), o usuário NÃO envia a cada uso.
MODELO_CROQUI = os.path.join(MODELOS_DIR, 'Modelo Croqui.xlsx')

MESES_PT = {
    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
    5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
    9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"
}

# Mapeamento coluna de entrada (0-indexed) -> célula do modelo de croqui
CROQUI_MAPPINGS = [
    (0, 'C42'),   # Coluna A -> OR
    (1, 'H53'),   # Coluna B -> TA
    (2, 'C53'),   # Coluna C -> OBRA (também usada no nome do arquivo)
    (3, 'S32'),   # Coluna D -> LOCAL
    (4, 'C51'),   # Coluna E -> CAUSA
    (5, 'B56'),   # Coluna F -> TRATATIVA
    (6, 'H31'),   # Coluna G -> ENDEREÇO
    (7, 'L43'),   # Coluna H -> EXEC
]


# =========================================================
# UI - CSS e JS compartilhados (inline, sem pasta static/)
# =========================================================
BASE_CSS = '''
:root {
    --navy: #0b2540; --navy-light: #123657; --blue: #1c6fb0; --blue-light: #2f8fd6;
    --mint: #8fd14f; --mint-dark: #6fb536; --ink: #1f2d3d; --muted: #64748b;
    --bg: #f4f7fb; --card: #ffffff; --border: #e2e8f0;
    --danger: #e2574c; --danger-bg: #fdeceb; --success: #2f9e63; --success-bg: #e8f7ee;
}
* { margin: 0; padding: 0; box-sizing: border-box; }
html { scroll-behavior: smooth; }
body { font-family: 'Inter', sans-serif; background: var(--bg); color: var(--ink); min-height: 100vh; }

/* ===== Topbar ===== */
.topbar {
    position: sticky; top: 0; left: 0; width: 100%; z-index: 100;
    background: linear-gradient(90deg, var(--navy) 0%, var(--navy-light) 100%);
    box-shadow: 0 2px 12px rgba(11, 37, 64, 0.15);
}
.topbar-inner {
    max-width: 1100px; margin: 0 auto; padding: 0 24px;
    display: flex; align-items: center; justify-content: space-between; height: 68px;
}
.brand { display: flex; align-items: center; gap: 12px; }
.brand-mark {
    width: 38px; height: 38px; border-radius: 11px; background: rgba(143, 209, 79, 0.15);
    color: var(--mint); display: flex; align-items: center; justify-content: center;
    font-size: 16px; flex-shrink: 0;
}
.brand-text { display: flex; flex-direction: column; line-height: 1.2; }
.brand-text strong { color: #ffffff; font-size: 14.5px; font-weight: 700; }
.brand-text span { color: rgba(255, 255, 255, 0.6); font-size: 12.5px; }

.topbar-nav { display: flex; align-items: center; gap: 6px; }
.nav-item {
    display: flex; align-items: center; gap: 9px; padding: 10px 16px; border-radius: 9px;
    color: rgba(255, 255, 255, 0.75); text-decoration: none; font-size: 14px; font-weight: 500;
    transition: background 0.2s ease, color 0.2s ease; position: relative; white-space: nowrap;
}
.nav-item i { font-size: 14px; }
.nav-item:hover { background: rgba(255, 255, 255, 0.08); color: #ffffff; }
.nav-item.active { background: rgba(143, 209, 79, 0.16); color: #ffffff; }
.nav-item.active i { color: var(--mint); }

.topbar-toggle {
    display: none; width: 40px; height: 40px; border-radius: 9px; border: none;
    background: rgba(255, 255, 255, 0.08); color: #fff; font-size: 16px; cursor: pointer;
    align-items: center; justify-content: center;
}

.content { max-width: 900px; margin: 0 auto; padding: 40px 24px 64px; }
.page-header { margin-bottom: 32px; text-align: center; }
.page-header p { margin-left: auto; margin-right: auto; }
.page-kicker {
    display: inline-flex; align-items: center; gap: 6px; color: var(--blue); font-size: 12.5px;
    font-weight: 600; letter-spacing: 0.04em; text-transform: uppercase; margin-bottom: 10px;
}
.page-header h1 { font-size: 28px; font-weight: 700; color: var(--navy); margin-bottom: 8px; }
.page-header p { color: var(--muted); font-size: 15px; max-width: 640px; line-height: 1.5; }

.card {
    background: var(--card); border: 1px solid var(--border); border-radius: 16px;
    margin-bottom: 24px; overflow: hidden; box-shadow: 0 1px 3px rgba(15, 23, 42, 0.04);
}
.card-header { background: linear-gradient(135deg, var(--blue) 0%, var(--navy) 100%); padding: 18px 26px; color: #fff; }
.card-header h2 { font-size: 16.5px; font-weight: 600; display: flex; align-items: center; gap: 9px; }
.card-header p { margin-top: 6px; font-size: 13px; color: rgba(255, 255, 255, 0.85); }
.card-body { padding: 26px; }

.columns-info { display: grid; grid-template-columns: repeat(4, 1fr); gap: 16px; margin-bottom: 22px; }
.column-card { background: var(--bg); border-radius: 12px; padding: 18px; border-left: 3px solid var(--blue); }
.column-number {
    background: var(--blue); color: #fff; width: 26px; height: 26px; border-radius: 50%;
    display: flex; align-items: center; justify-content: center; margin-bottom: 10px;
    font-weight: 700; font-size: 12.5px;
}
.column-title { font-size: 14.5px; font-weight: 700; color: var(--navy); margin-bottom: 6px; }
.column-desc { font-size: 12.5px; color: var(--muted); line-height: 1.45; }

.info-note {
    background: var(--bg); border-left: 3px solid var(--mint-dark); padding: 13px 16px; border-radius: 8px;
    display: flex; align-items: flex-start; gap: 11px; font-size: 13.5px; color: #3d4a5c; line-height: 1.5;
}
.info-note i { color: var(--mint-dark); margin-top: 2px; }
.warning-note {
    background: #fff8e6; border-left: 3px solid #e0a721; padding: 13px 16px; border-radius: 8px;
    margin-top: 16px; display: flex; align-items: flex-start; gap: 11px; font-size: 13.5px;
    color: #7a5c0e; line-height: 1.5;
}
.warning-note i { color: #e0a721; margin-top: 2px; }

.upload-area {
    border: 2px dashed #c7d5e6; border-radius: 14px; padding: 40px 26px; text-align: center;
    cursor: pointer; transition: all 0.25s ease; background: #f9fbfd; margin-bottom: 22px;
}
.upload-area:hover, .upload-area.dragover { border-color: var(--blue-light); background: #eef6fc; }
.upload-icon { font-size: 44px; color: var(--blue); margin-bottom: 14px; }
.upload-text { font-size: 16.5px; font-weight: 600; color: var(--navy); margin-bottom: 6px; }
.upload-subtext { font-size: 13.5px; color: var(--muted); }

.file-info {
    display: none; background: #f2f6fa; border-radius: 10px; padding: 13px 16px; margin-top: -6px;
    margin-bottom: 22px; border: 1px solid var(--border); align-items: center; justify-content: space-between;
}
.file-info.show { display: flex; }
.file-name { font-weight: 500; color: var(--navy); display: flex; align-items: center; gap: 10px; font-size: 13.5px; }
.file-size { color: var(--muted); font-size: 12.5px; }
.remove-file { background: none; border: none; color: var(--danger); cursor: pointer; font-size: 18px; }

.btn-submit {
    width: 100%; padding: 15px; background: linear-gradient(135deg, var(--mint) 0%, var(--mint-dark) 100%);
    color: #143a0e; border: none; border-radius: 10px; font-size: 15.5px; font-weight: 700; cursor: pointer;
    display: flex; align-items: center; justify-content: center; gap: 9px;
    transition: transform 0.2s ease, box-shadow 0.2s ease;
}
.btn-submit:hover:not(:disabled) { transform: translateY(-2px); box-shadow: 0 10px 24px rgba(111, 181, 54, 0.35); }
.btn-submit:disabled { opacity: 0.5; cursor: not-allowed; }

.loading { display: none; text-align: center; padding: 22px; }
.loading.show { display: block; }
.spinner {
    border: 4px solid #e2e8f0; border-top: 4px solid var(--blue); border-radius: 50%;
    width: 42px; height: 42px; animation: spin 1s linear infinite; margin: 0 auto 16px;
}
@keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }

.error-message, .success-message { padding: 12px 16px; border-radius: 8px; margin-top: 14px; display: none; font-size: 13.5px; }
.error-message { background: var(--danger-bg); border: 1px solid #f2b8b3; color: #a4342a; }
.success-message { background: var(--success-bg); border: 1px solid #a6dfbd; color: #1f7a48; }
.error-message.show, .success-message.show { display: block; }

@media (max-width: 720px) {
    .topbar-nav {
        position: absolute; top: 68px; left: 0; right: 0; flex-direction: column; align-items: stretch;
        background: var(--navy-light); padding: 8px 16px 16px; gap: 4px;
        box-shadow: 0 8px 16px rgba(11, 37, 64, 0.2);
        display: none;
    }
    .topbar-nav.open { display: flex; }
    .nav-item { padding: 12px 14px; }
    .topbar-toggle { display: flex; }
    .content { padding: 28px 18px 48px; }
    .columns-info { grid-template-columns: 1fr; }
    .card-body { padding: 20px; }
}

@media (min-width: 721px) and (max-width: 980px) {
    .columns-info { grid-template-columns: repeat(2, 1fr); }
}
'''

UPLOAD_JS = '''
(function () {
    var endpoint = "__ENDPOINT__";
    var defaultZipName = "__ZIPNAME__";

    var dropZone = document.getElementById('dropZone');
    var fileInput = document.getElementById('fileInput');
    var fileInfo = document.getElementById('fileInfo');
    var fileName = document.getElementById('fileName');
    var fileSize = document.getElementById('fileSize');
    var removeFile = document.getElementById('removeFile');
    var submitBtn = document.getElementById('submitBtn');
    var uploadForm = document.getElementById('uploadForm');
    var loading = document.getElementById('loading');
    var errorMessage = document.getElementById('errorMessage');
    var errorText = document.getElementById('errorText');
    var successMessage = document.getElementById('successMessage');
    var successText = document.getElementById('successText');

    var selectedFile = null;

    dropZone.addEventListener('click', function () { fileInput.click(); });
    fileInput.addEventListener('change', function (e) { handleFile(e.target.files[0]); });

    dropZone.addEventListener('dragover', function (e) { e.preventDefault(); dropZone.classList.add('dragover'); });
    dropZone.addEventListener('dragleave', function () { dropZone.classList.remove('dragover'); });
    dropZone.addEventListener('drop', function (e) {
        e.preventDefault();
        dropZone.classList.remove('dragover');
        handleFile(e.dataTransfer.files[0]);
    });

    removeFile.addEventListener('click', function (e) {
        e.stopPropagation();
        selectedFile = null;
        fileInput.value = '';
        fileInfo.classList.remove('show');
        submitBtn.disabled = true;
    });

    function handleFile(file) {
        if (!file) return;
        var validTypes = [
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.ms-excel'
        ];
        if (validTypes.indexOf(file.type) === -1 && !file.name.match(/\\.(xlsx|xls)$/i)) {
            showError('Por favor, selecione um arquivo Excel válido (.xlsx ou .xls)');
            return;
        }
        selectedFile = file;
        fileName.textContent = file.name;
        fileSize.textContent = formatFileSize(file.size);
        fileInfo.classList.add('show');
        submitBtn.disabled = false;
        hideMessages();
    }

    function formatFileSize(bytes) {
        if (bytes === 0) return '0 Bytes';
        var k = 1024;
        var sizes = ['Bytes', 'KB', 'MB', 'GB'];
        var i = Math.floor(Math.log(bytes) / Math.log(k));
        return parseFloat((bytes / Math.pow(k, i)).toFixed(2)) + ' ' + sizes[i];
    }

    uploadForm.addEventListener('submit', function (e) {
        e.preventDefault();
        if (!selectedFile) {
            showError('Por favor, selecione um arquivo primeiro');
            return;
        }

        loading.classList.add('show');
        submitBtn.disabled = true;
        hideMessages();

        var formData = new FormData();
        formData.append('planilha', selectedFile);

        fetch(endpoint, { method: 'POST', body: formData })
            .then(function (response) {
                if (!response.ok) {
                    return response.json().then(function (errorData) {
                        throw new Error(errorData.error || 'Erro ao processar arquivo');
                    });
                }
                var contentDisposition = response.headers.get('Content-Disposition');
                var filename = defaultZipName;
                if (contentDisposition) {
                    var filenameMatch = contentDisposition.match(/filename="?([^"]+)"?/);
                    if (filenameMatch) filename = filenameMatch[1];
                }
                return response.blob().then(function (blob) { return { blob: blob, filename: filename }; });
            })
            .then(function (result) {
                var url = window.URL.createObjectURL(result.blob);
                var a = document.createElement('a');
                a.href = url;
                a.download = result.filename;
                document.body.appendChild(a);
                a.click();
                window.URL.revokeObjectURL(url);
                a.remove();

                showSuccess('Arquivo processado com sucesso! O download começará automaticamente.');

                setTimeout(function () {
                    selectedFile = null;
                    fileInput.value = '';
                    fileInfo.classList.remove('show');
                    submitBtn.disabled = true;
                }, 2000);
            })
            .catch(function (error) {
                showError(error.message);
            })
            .finally(function () {
                loading.classList.remove('show');
                submitBtn.disabled = false;
            });
    });

    function showError(message) {
        errorText.textContent = message;
        errorMessage.classList.add('show');
        successMessage.classList.remove('show');
    }
    function showSuccess(message) {
        successText.textContent = message;
        successMessage.classList.add('show');
        errorMessage.classList.remove('show');
    }
    function hideMessages() {
        errorMessage.classList.remove('show');
        successMessage.classList.remove('show');
    }
})();
'''

PAGE_TEMPLATE = '''
<!DOCTYPE html>
<html lang="pt-br">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{{ title }}</title>
<link rel="icon" type="image/svg+xml" href="data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHZpZXdCb3g9IjAgMCA2NCA2NCI+CjxyZWN0IHdpZHRoPSI2NCIgaGVpZ2h0PSI2NCIgcng9IjE0IiBmaWxsPSIjMGIyNTQwIi8+Cjx0ZXh0IHg9IjMyIiB5PSI0MyIgZm9udC1mYW1pbHk9IkFyaWFsLCBIZWx2ZXRpY2EsIHNhbnMtc2VyaWYiIGZvbnQtc2l6ZT0iMjYiIGZvbnQtd2VpZ2h0PSI3MDAiIGZpbGw9IiM4ZmQxNGYiIHRleHQtYW5jaG9yPSJtaWRkbGUiPlNEPC90ZXh0Pgo8L3N2Zz4=">
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap" rel="stylesheet">
<link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
<style>{{ base_css|safe }}</style>
</head>
<body>

<header class="topbar">
    <div class="topbar-inner">
        <div class="brand">
            <div class="brand-mark"><i class="fas fa-tower-broadcast"></i></div>
            <div class="brand-text"><strong>Sistema de</strong><span>Documentos</span></div>
        </div>

        <nav class="topbar-nav" id="topbarNav">
            <a href="{{ url_for('pleito') }}" class="nav-item {{ 'active' if active == 'pleito' else '' }}">
                <i class="fas fa-file-word"></i><span>Carta Pleito</span>
            </a>
            <a href="{{ url_for('croqui') }}" class="nav-item {{ 'active' if active == 'croqui' else '' }}">
                <i class="fas fa-draw-polygon"></i><span>Criação de Croqui</span>
            </a>
        </nav>

        <button class="topbar-toggle" id="topbarToggle" aria-label="Abrir menu"><i class="fas fa-bars"></i></button>
    </div>
</header>

<main class="content">
    <div class="page-header">
        <div class="page-kicker"><i class="{{ kicker_icon }}"></i> Modelos</div>
        <h1>{{ heading }}</h1>
        <p>{{ description }}</p>
    </div>

    {{ body_html|safe }}
</main>

<script>
    document.getElementById('topbarToggle').addEventListener('click', function () {
        document.getElementById('topbarNav').classList.toggle('open');
    });
</script>
<script>{{ upload_js|safe }}</script>
</body>
</html>
'''

PLEITO_BODY = '''
<div class="card">
    <div class="card-header">
        <h2><i class="fas fa-circle-info"></i> Instruções para o Arquivo Excel</h2>
        <p>Para garantir o correto processamento, seu arquivo deve seguir o formato abaixo</p>
    </div>
    <div class="card-body">
        <div class="columns-info">
            <div class="column-card"><div class="column-number">1</div><div class="column-title">Tipo</div><div class="column-desc">Nacional ou Regional</div></div>
            <div class="column-card"><div class="column-number">2</div><div class="column-title">Sequência</div><div class="column-desc">Número da TA</div></div>
            <div class="column-card"><div class="column-number">3</div><div class="column-title">Data de Criação</div><div class="column-desc">Data da ocorrência</div></div>
            <div class="column-card"><div class="column-number">4</div><div class="column-title">Município</div><div class="column-desc">Localidade da ocorrência</div></div>
        </div>
        <div class="warning-note">
            <i class="fas fa-triangle-exclamation"></i>
            <span><strong>Importante:</strong> as colunas devem estar exatamente nesta ordem: Tipo, Sequência, Data de Criação e Município. A primeira linha pode conter cabeçalhos.</span>
        </div>
    </div>
</div>

<div class="card">
    <div class="card-header">
        <h2><i class="fas fa-cloud-arrow-up"></i> Enviar planilha</h2>
        <p>Formatos aceitos: .xlsx, .xls</p>
    </div>
    <div class="card-body">
        <form id="uploadForm" enctype="multipart/form-data">
            <div class="upload-area" id="dropZone">
                <div class="upload-icon"><i class="fas fa-cloud-arrow-up"></i></div>
                <div class="upload-text">Arraste e solte sua planilha aqui</div>
                <div class="upload-subtext">ou clique para selecionar</div>
                <input type="file" id="fileInput" name="planilha" accept=".xlsx,.xls" style="display:none;">
            </div>
            <div class="file-info" id="fileInfo">
                <div class="file-name"><i class="fas fa-file-excel" style="color:#2f9e63;font-size:20px;"></i><span id="fileName"></span></div>
                <div><span class="file-size" id="fileSize"></span><button type="button" class="remove-file" id="removeFile"><i class="fas fa-circle-xmark"></i></button></div>
            </div>
            <button type="submit" class="btn-submit" id="submitBtn" disabled><i class="fas fa-paper-plane"></i> Processar e Gerar Cartas</button>
        </form>
        <div class="loading" id="loading">
            <div class="spinner"></div>
            <p style="color:var(--blue);font-weight:600;">Processando arquivo...</p>
            <p style="color:var(--muted);font-size:13px;">Isso pode levar alguns segundos</p>
        </div>
        <div class="error-message" id="errorMessage"><i class="fas fa-circle-exclamation"></i> <span id="errorText"></span></div>
        <div class="success-message" id="successMessage"><i class="fas fa-circle-check"></i> <span id="successText"></span></div>
    </div>
</div>
'''

CROQUI_BODY = '''
<div class="card">
    <div class="card-header">
        <h2><i class="fas fa-circle-info"></i> Instruções para a Planilha de Entrada</h2>
        <p>As colunas devem seguir exatamente esta ordem, a partir da coluna A</p>
    </div>
    <div class="card-body">
        <div class="columns-info">
            <div class="column-card"><div class="column-number">A</div><div class="column-title">OR</div><div class="column-desc">Ordem de Requisição</div></div>
            <div class="column-card"><div class="column-number">B</div><div class="column-title">TA</div><div class="column-desc">Número da TA</div></div>
            <div class="column-card"><div class="column-number">C</div><div class="column-title">Obra</div><div class="column-desc">Também usada para nomear o arquivo gerado</div></div>
            <div class="column-card"><div class="column-number">D</div><div class="column-title">Local</div><div class="column-desc">Localidade da obra</div></div>
            <div class="column-card"><div class="column-number">E</div><div class="column-title">Causa</div><div class="column-desc">Motivo da ocorrência</div></div>
            <div class="column-card"><div class="column-number">F</div><div class="column-title">Tratativa</div><div class="column-desc">Ação tomada</div></div>
            <div class="column-card"><div class="column-number">G</div><div class="column-title">Endereço</div><div class="column-desc">Endereço completo</div></div>
            <div class="column-card"><div class="column-number">H</div><div class="column-title">Exec</div><div class="column-desc">Responsável pela execução</div></div>
        </div>
        <div class="info-note">
            <i class="fas fa-circle-check"></i>
            <span>O modelo de croqui utilizado é fixo e já está configurado no sistema pela equipe. Se o modelo precisar ser atualizado, fale com o administrador do sistema.</span>
        </div>
        <div class="warning-note">
            <i class="fas fa-triangle-exclamation"></i>
            <span><strong>Importante:</strong> a primeira linha pode conter cabeçalhos. Linhas totalmente vazias nas 8 primeiras colunas são ignoradas.</span>
        </div>
    </div>
</div>

<div class="card">
    <div class="card-header">
        <h2><i class="fas fa-cloud-arrow-up"></i> Enviar planilha de entrada</h2>
        <p>Formatos aceitos: .xlsx, .xls</p>
    </div>
    <div class="card-body">
        <form id="uploadForm" enctype="multipart/form-data">
            <div class="upload-area" id="dropZone">
                <div class="upload-icon"><i class="fas fa-cloud-arrow-up"></i></div>
                <div class="upload-text">Arraste e solte a planilha de entrada aqui</div>
                <div class="upload-subtext">ou clique para selecionar</div>
                <input type="file" id="fileInput" name="planilha" accept=".xlsx,.xls" style="display:none;">
            </div>
            <div class="file-info" id="fileInfo">
                <div class="file-name"><i class="fas fa-file-excel" style="color:#2f9e63;font-size:20px;"></i><span id="fileName"></span></div>
                <div><span class="file-size" id="fileSize"></span><button type="button" class="remove-file" id="removeFile"><i class="fas fa-circle-xmark"></i></button></div>
            </div>
            <button type="submit" class="btn-submit" id="submitBtn" disabled><i class="fas fa-paper-plane"></i> Processar e Gerar Croquis</button>
        </form>
        <div class="loading" id="loading">
            <div class="spinner"></div>
            <p style="color:var(--blue);font-weight:600;">Gerando croquis...</p>
            <p style="color:var(--muted);font-size:13px;">Isso pode levar alguns segundos</p>
        </div>
        <div class="error-message" id="errorMessage"><i class="fas fa-circle-exclamation"></i> <span id="errorText"></span></div>
        <div class="success-message" id="successMessage"><i class="fas fa-circle-check"></i> <span id="successText"></span></div>
    </div>
</div>
'''


def render_page(active, title, kicker_icon, heading, description, body_html, endpoint, zip_name):
    js = UPLOAD_JS.replace('__ENDPOINT__', endpoint).replace('__ZIPNAME__', zip_name)
    return render_template_string(
        PAGE_TEMPLATE,
        active=active, title=title, kicker_icon=kicker_icon, heading=heading,
        description=description, body_html=body_html, base_css=BASE_CSS, upload_js=js
    )


# =========================================================
# Rotas de página
# =========================================================
@app.route('/')
def index():
    return redirect(url_for('pleito'))


@app.route('/pleito')
def pleito():
    return render_page(
        active='pleito',
        title='Carta Pleito · Sistema de Documentos',
        kicker_icon='fas fa-file-word',
        heading='Carta Pleito',
        description='Envie a planilha com as ocorrências e gere automaticamente as cartas consolidadas (Nacional e Regional) e a planilha de acompanhamento formatada.',
        body_html=PLEITO_BODY,
        endpoint='/processar',
        zip_name='cartas_pleito.zip'
    )


@app.route('/croqui')
def croqui():
    return render_page(
        active='croqui',
        title='Criação de Croqui · Sistema de Documentos',
        kicker_icon='fas fa-draw-polygon',
        heading='Criação de Croqui',
        description='Envie a planilha de entrada com os dados das obras. O modelo de croqui já está cadastrado no sistema — não é necessário enviá-lo.',
        body_html=CROQUI_BODY,
        endpoint='/processar-croqui',
        zip_name='croquis_gerados.zip'
    )


# =========================================================
# Helpers - Carta Pleito (lógica original, inalterada)
# =========================================================
def apply_line_spacing(paragraph):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def substituir_placeholders(text, numero_ta, data_ocorrencia, regiao, tipo, data_formatada):
    if 'Nº DA OBRAS/SERVIÇOS: Nº TA' in text:
        text = text.replace('Nº DA OBRAS/SERVIÇOS: Nº TA', f'Nº DA OBRAS/SERVIÇOS: Nº TA {numero_ta}')
    if 'OCORRÊNCIA Nº TA' in text:
        text = text.replace('OCORRÊNCIA Nº TA', f'OCORRÊNCIA Nº TA {numero_ta}')
    if 'DATA DA CONSTATAÇÃO DA OCORRÊNCIA:' in text:
        text = text.replace('DATA DA CONSTATAÇÃO DA OCORRÊNCIA:', f'DATA DA CONSTATAÇÃO DA OCORRÊNCIA: {data_ocorrencia}')
    if tipo == 'regional':
        if 'Em referência à TA' in text:
            text = text.replace('Em referência à TA', f'Em referência à TA {numero_ta}')
    else:
        if 'Recebemos em nosso sistema a TA' in text:
            text = text.replace('TA', f'TA {numero_ta}')
    if 'REGIÃO:' in text:
        partes = text.split('REGIÃO:')
        text = f'{partes[0]}REGIÃO: {regiao}'
    if 'Data assinatura coordenador:' in text:
        text = text.replace('Data assinatura coordenador:', f'Data assinatura coordenador: {data_formatada}')
    return text


def add_carta_to_document(doc, template_path, numero_ta, data_ocorrencia, regiao, tipo, data_formatada, is_first=False):
    if not is_first:
        doc.add_page_break()

    modelo_doc = Document(template_path)

    for element in modelo_doc.element.body:
        tag = element.tag.split('}')[-1]

        if tag == 'p':
            paragraph = next((p for p in modelo_doc.paragraphs if p._element is element), None)
            if paragraph is None:
                continue

            new_paragraph = doc.add_paragraph()
            new_paragraph.style = paragraph.style

            if len(paragraph.runs) <= 1:
                text = substituir_placeholders(paragraph.text, numero_ta, data_ocorrencia, regiao, tipo, data_formatada)
                new_paragraph.add_run(text)
            else:
                for run in paragraph.runs:
                    run_text = substituir_placeholders(run.text, numero_ta, data_ocorrencia, regiao, tipo, data_formatada)
                    new_run = new_paragraph.add_run(run_text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline

            apply_line_spacing(new_paragraph)

        elif tag == 'tbl':
            table = next((t for t in modelo_doc.tables if t._element is element), None)
            if table is None:
                continue

            new_table = doc.add_table(rows=len(table.rows), cols=len(table.columns))
            new_table.style = table.style
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_table.cell(i, j).text = substituir_placeholders(
                        cell.text, numero_ta, data_ocorrencia, regiao, tipo, data_formatada
                    )


def formatar_excel(excel_path, df):
    wb = Workbook()
    ws = wb.active
    ws.title = "Pleitos"

    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    header_font = Font(name='Calibri', size=11, bold=True, color="FFFFFF")

    even_row_fill = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
    odd_row_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

    cell_font = Font(name='Calibri', size=10)
    cell_alignment = Alignment(horizontal='center', vertical='center')

    thin_border = Border(
        left=Side(style='thin', color='B4C6E7'),
        right=Side(style='thin', color='B4C6E7'),
        top=Side(style='thin', color='B4C6E7'),
        bottom=Side(style='thin', color='B4C6E7')
    )

    headers = list(df.columns)
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = thin_border

    for row_idx, row in df.iterrows():
        fill = even_row_fill if (row_idx % 2 == 0) else odd_row_fill
        for col_idx, value in enumerate(row, 1):
            cell = ws.cell(row=row_idx + 2, column=col_idx, value=value)
            cell.font = cell_font
            cell.alignment = cell_alignment
            cell.border = thin_border
            cell.fill = fill

    for col_idx in range(1, len(headers) + 1):
        max_length = 0
        column_letter = get_column_letter(col_idx)
        for row in ws.iter_rows(min_col=col_idx, max_col=col_idx):
            for cell in row:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[column_letter].width = min(max_length + 3, 50)

    ws.freeze_panes = 'A2'
    ws.auto_filter.ref = ws.dimensions
    wb.save(excel_path)


# =========================================================
# Rota de processamento - Carta Pleito (lógica original, inalterada)
# =========================================================
@app.route('/processar', methods=['POST'])
def processar():
    if not os.path.exists(MODELOS_DIR):
        return jsonify({'error': 'Diretório de modelos não encontrado'}), 500

    if not os.path.exists(MODELO_NACIONAL) or not os.path.exists(MODELO_REGIONAL):
        return jsonify({'error': 'Modelos de carta não encontrados no servidor'}), 500

    if 'planilha' not in request.files:
        return jsonify({'error': 'Nenhum arquivo enviado'}), 400

    file = request.files['planilha']
    if file.filename == '':
        return jsonify({'error': 'Nome de arquivo inválido'}), 400

    try:
        hoje = datetime.now()
        data_formatada = hoje.strftime("%d/%m/%Y")
        mes_extenso = MESES_PT[hoje.month]
        data_completa = hoje.strftime("%d-%m-%Y")

        df = pd.read_excel(file)

        if df.empty:
            return jsonify({'error': 'A planilha está vazia'}), 400

        if len(df.columns) < 4:
            return jsonify({'error': 'A planilha deve ter pelo menos 4 colunas (Tipo, TA, Data, Local)'}), 400

        col_tipo, col_ta, col_data, col_local = df.columns[:4]

        temp_dir = tempfile.mkdtemp()
        arquivos_gerados = []
        dados_saida = []

        for _, row in df.iterrows():
            tipo = str(row[col_tipo]).strip().lower()
            ta = str(row[col_ta])
            localidade = str(row[col_local])

            if tipo not in ['nacional', 'regional']:
                continue

            dados_saida.append({
                'UF': 'MA',
                'TA': ta,
                'Pleito': 'Analisando',
                'Aprovador': 'Andressa Alves Vieira Boaventura' if tipo == 'nacional' else 'André Luiz Ferreira de Carvalho',
                'Seguimento': 'Nacional' if tipo == 'nacional' else 'Regional',
                'Mes': mes_extenso,
                'Responsável': 'Orlando/Anderson',
                'Localidade': localidade,
                'Observações': ''
            })

        if any(d['Seguimento'] == 'Nacional' for d in dados_saida):
            doc_nac = Document(MODELO_NACIONAL)
            for element in list(doc_nac.element.body):
                if element.tag.endswith('p') or element.tag.endswith('tbl'):
                    doc_nac.element.body.remove(element)

            first = True
            for item in dados_saida:
                if item['Seguimento'] == 'Nacional':
                    add_carta_to_document(doc_nac, MODELO_NACIONAL, item['TA'], '', item['Localidade'], 'nacional', data_formatada, first)
                    first = False

            if not first:
                output_nac = os.path.join(temp_dir, f'cartas_pleito_nacional_consolidado_{data_completa}.docx')
                doc_nac.save(output_nac)
                arquivos_gerados.append(output_nac)

        if any(d['Seguimento'] == 'Regional' for d in dados_saida):
            doc_reg = Document(MODELO_REGIONAL)
            for element in list(doc_reg.element.body):
                if element.tag.endswith('p') or element.tag.endswith('tbl'):
                    doc_reg.element.body.remove(element)

            first = True
            for item in dados_saida:
                if item['Seguimento'] == 'Regional':
                    add_carta_to_document(doc_reg, MODELO_REGIONAL, item['TA'], '', item['Localidade'], 'regional', data_formatada, first)
                    first = False

            if not first:
                output_reg = os.path.join(temp_dir, f'cartas_pleito_regional_consolidado_{data_completa}.docx')
                doc_reg.save(output_reg)
                arquivos_gerados.append(output_reg)

        if dados_saida:
            output_excel = os.path.join(temp_dir, f'CartasPleitos({data_completa}).xlsx')
            formatar_excel(output_excel, pd.DataFrame(dados_saida))
            arquivos_gerados.append(output_excel)

        if not arquivos_gerados:
            return jsonify({'error': 'Nenhum dado válido para processar'}), 400

        zip_buffer = io.BytesIO()
        with ZipFile(zip_buffer, 'w') as zipf:
            for arquivo in arquivos_gerados:
                zipf.write(arquivo, os.path.basename(arquivo))
        zip_buffer.seek(0)

        shutil.rmtree(temp_dir, ignore_errors=True)

        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'cartas_pleito_{data_completa}.zip'
        )

    except Exception as e:
        return jsonify({'error': f'Erro ao processar arquivo: {str(e)}'}), 500


# =========================================================
# Preenchimento de células preservando imagens/estilos do modelo
# =========================================================
# O openpyxl documenta oficialmente que abrir e salvar um .xlsx com
# load_workbook()/save() descarta imagens e gráficos do arquivo original.
# Além disso, reescrever o XML inteiro com parsers genéricos (ElementTree
# incluído) troca os prefixos de namespaces estendidos que o Excel real
# usa (x14ac, xr, xr2, xr3, mc:Ignorable), o que faz o Excel acusar
# "conteúdo com problema" ao reabrir o arquivo gerado.
#
# Por isso, em vez de fazer parse/reserialização da árvore XML inteira,
# a função abaixo faz uma edição cirúrgica de texto: localiza a tag <c>
# exata de cada célula-alvo dentro do XML da planilha ativa e substitui
# só o conteúdo dela. Todo o resto do arquivo — o restante do próprio
# XML, imagens, desenhos, estilos, namespaces estendidos — permanece
# byte a byte idêntico ao modelo original.

def _col_to_num(col_letters):
    num = 0
    for ch in col_letters:
        num = num * 26 + (ord(ch) - ord('A') + 1)
    return num


def _split_ref(ref):
    m = re.match(r'^([A-Z]+)(\d+)$', ref)
    return m.group(1), int(m.group(2))


def _find_active_sheet_path(zf):
    import re as _re
    wb_xml = zf.read('xl/workbook.xml').decode('utf-8')
    rels_xml = zf.read('xl/_rels/workbook.xml.rels').decode('utf-8')

    active_index = 0
    m = _re.search(r'<workbookView[^>]*activeTab="(\d+)"', wb_xml)
    if m:
        active_index = int(m.group(1))

    sheet_matches = _re.findall(r'<sheet\b[^>]*/>', wb_xml)
    if not sheet_matches or active_index >= len(sheet_matches):
        active_index = 0
    target_sheet_tag = sheet_matches[active_index]
    rid_match = _re.search(r'r:id="([^"]+)"', target_sheet_tag)
    target_rid = rid_match.group(1)

    rel_map = {}
    for tag in _re.findall(r'<Relationship\b[^>]*/>', rels_xml):
        id_m = _re.search(r'\sId="([^"]+)"', tag)
        target_m = _re.search(r'\sTarget="([^"]+)"', tag)
        if id_m and target_m:
            rel_map[id_m.group(1)] = target_m.group(1)

    target = rel_map[target_rid]
    target = target.lstrip('/')
    if not target.startswith('xl/'):
        target = 'xl/' + target
    return target


def _build_merge_map(sheet_xml):
    merge_map = {}
    for m in re.finditer(r'<mergeCell\s+ref="([A-Z]+\d+):([A-Z]+\d+)"\s*/>', sheet_xml):
        start, end = m.group(1), m.group(2)
        start_col, start_row = _split_ref(start)
        end_col, end_row = _split_ref(end)
        c1, c2 = _col_to_num(start_col), _col_to_num(end_col)
        for r in range(start_row, end_row + 1):
            for c in range(min(c1, c2), max(c1, c2) + 1):
                col_letters = ''
                n = c
                while n > 0:
                    n, rem = divmod(n - 1, 26)
                    col_letters = chr(65 + rem) + col_letters
                merge_map[f'{col_letters}{r}'] = start
    return merge_map


def _escape_xml_text(text):
    return (
        text.replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
    )


def _build_cell_xml(cell_ref, style_attr, value):
    if isinstance(value, bool):
        return f'<c r="{cell_ref}"{style_attr} t="b"><v>{"1" if value else "0"}</v></c>'
    if isinstance(value, (int, float)):
        v_text = repr(value) if isinstance(value, float) else str(value)
        return f'<c r="{cell_ref}"{style_attr}><v>{v_text}</v></c>'
    text = _escape_xml_text('' if value is None else str(value))
    return f'<c r="{cell_ref}"{style_attr} t="inlineStr"><is><t xml:space="preserve">{text}</t></is></c>'


_CELL_PATTERN_TMPL = r'<c r="{ref}"(?P<attrs>[^>]*?)(?:/>|>(?P<content>.*?)</c>)'


def _extract_style_attr(attrs_str):
    m = re.search(r'\ss="(\d+)"', attrs_str)
    return f' s="{m.group(1)}"' if m else ''


def _set_cell_in_row(row_xml, cell_ref, new_cell_xml):
    """Substitui a célula se existir; caso contrário insere na posição correta dentro da linha."""
    pattern = re.compile(_CELL_PATTERN_TMPL.format(ref=re.escape(cell_ref)), re.DOTALL)
    match = pattern.search(row_xml)
    if match:
        return row_xml[:match.start()] + new_cell_xml + row_xml[match.end():]

    # célula não existe nessa linha: insere na posição correta (ordem crescente de coluna)
    target_col_num = _col_to_num(_split_ref(cell_ref)[0])
    insert_pos = None
    for m in re.finditer(r'<c r="([A-Z]+)\d+"[^>]*?(?:/>|>.*?</c>)', row_xml, re.DOTALL):
        existing_col_num = _col_to_num(m.group(1))
        if existing_col_num > target_col_num:
            insert_pos = m.start()
            break
    if insert_pos is None:
        # nenhuma célula com coluna maior: insere antes de </row>
        close_idx = row_xml.rindex('</row>')
        return row_xml[:close_idx] + new_cell_xml + row_xml[close_idx:]
    return row_xml[:insert_pos] + new_cell_xml + row_xml[insert_pos:]


def _set_cell_value(sheet_xml, cell_ref, value):
    col_letters, row_num = _split_ref(cell_ref)

    row_pattern = re.compile(
        r'(<row r="' + str(row_num) + r'"[^>]*?)(/>|>(.*?)</row>)', re.DOTALL
    )
    match = row_pattern.search(sheet_xml)

    if match:
        row_open_attrs = match.group(1)
        is_self_closing = match.group(2).startswith('/>')
        row_content = '' if is_self_closing else (match.group(3) or '')

        existing_style = ''
        cell_match = re.search(
            _CELL_PATTERN_TMPL.format(ref=re.escape(cell_ref)), row_content, re.DOTALL
        )
        if cell_match:
            existing_style = _extract_style_attr(cell_match.group('attrs'))

        new_cell_xml = _build_cell_xml(cell_ref, existing_style, value)
        new_row_content = _set_cell_in_row(row_content, cell_ref, new_cell_xml)
        new_row_xml = f'{row_open_attrs}>{new_row_content}</row>'

        return sheet_xml[:match.start()] + new_row_xml + sheet_xml[match.end():]

    # a linha inteira não existe: cria uma nova <row> na posição correta dentro de <sheetData>
    new_cell_xml = _build_cell_xml(cell_ref, '', value)
    new_row_xml = f'<row r="{row_num}">{new_cell_xml}</row>'

    sheet_data_match = re.search(r'<sheetData>(.*?)</sheetData>', sheet_xml, re.DOTALL)
    sheet_data_content = sheet_data_match.group(1)

    insert_pos_in_data = None
    for m in re.finditer(r'<row r="(\d+)"[^>]*?(?:/>|>.*?</row>)', sheet_data_content, re.DOTALL):
        if int(m.group(1)) > row_num:
            insert_pos_in_data = m.start()
            break

    if insert_pos_in_data is None:
        new_sheet_data_content = sheet_data_content + new_row_xml
    else:
        new_sheet_data_content = (
            sheet_data_content[:insert_pos_in_data] + new_row_xml + sheet_data_content[insert_pos_in_data:]
        )

    start, end = sheet_data_match.span(1)
    return sheet_xml[:start] + new_sheet_data_content + sheet_xml[end:]


def _safe_date_time(date_time):
    year, month, day, hour, minute, second = date_time
    year = max(year, 1980)
    month = min(max(month, 1), 12)
    day = min(max(day, 1), 31)
    hour = min(max(hour, 0), 23)
    minute = min(max(minute, 0), 59)
    second = min(max(second, 0), 59)
    return (year, month, day, hour, minute, second)


def preencher_modelo_xlsx(modelo_path, valores, output_path):
    """
    Preenche células de um .xlsx a partir de {celula: valor} editando SOMENTE
    o texto das tags <c> alvo dentro do XML da planilha ativa — nunca faz o
    parse/reserialização da árvore inteira. Isso evita o problema clássico de
    bibliotecas XML (incluindo ElementTree e openpyxl) trocarem prefixos de
    namespace (mc:Ignorable, x14ac, xr, xr2, xr3) usados por arquivos gerados
    pelo Excel de verdade, o que faz o Excel acusar "conteúdo com problema"
    ao reabrir. Todo o resto do arquivo — incluindo o restante do próprio
    XML da planilha, imagens, desenhos, estilos — permanece byte a byte
    idêntico ao modelo original.
    """
    with zipfile.ZipFile(modelo_path, 'r') as zin:
        sheet_path = _find_active_sheet_path(zin)
        sheet_xml = zin.read(sheet_path).decode('utf-8')

        merge_map = _build_merge_map(sheet_xml)

        for cell_ref, value in valores.items():
            if value is None:
                continue
            actual_ref = merge_map.get(cell_ref, cell_ref)
            sheet_xml = _set_cell_value(sheet_xml, actual_ref, value)

        new_sheet_bytes = sheet_xml.encode('utf-8')

        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = new_sheet_bytes if item.filename == sheet_path else zin.read(item.filename)
                fresh_info = zipfile.ZipInfo(
                    filename=item.filename,
                    date_time=_safe_date_time(item.date_time)
                )
                fresh_info.compress_type = zipfile.ZIP_DEFLATED
                fresh_info.external_attr = item.external_attr
                fresh_info.create_system = item.create_system
                zout.writestr(fresh_info, data)

        with zipfile.ZipFile(output_path, 'r') as check:
            bad_file = check.testzip()
            if bad_file is not None:
                raise RuntimeError(f'Falha de integridade ao gerar o croqui (entrada corrompida: {bad_file})')




# =========================================================
# Helpers - Criação de Croqui
# =========================================================
def sanitize_filename(filename):
    invalid_chars = r'[<>:"/\\|?*]'
    sanitized = re.sub(invalid_chars, '_', str(filename))
    return sanitized.strip()[:100]


# =========================================================
# Rota de processamento - Criação de Croqui
# =========================================================
@app.route('/processar-croqui', methods=['POST'])
def processar_croqui():
    if not os.path.exists(MODELO_CROQUI):
        return jsonify({'error': 'Modelo de croqui não encontrado no servidor. Contate o administrador.'}), 500

    if 'planilha' not in request.files:
        return jsonify({'error': 'Nenhum arquivo enviado'}), 400

    file = request.files['planilha']
    if file.filename == '':
        return jsonify({'error': 'Nome de arquivo inválido'}), 400

    temp_dir = tempfile.mkdtemp()

    try:
        input_path = os.path.join(temp_dir, 'entrada.xlsx')
        file.save(input_path)

        wb_input = openpyxl.load_workbook(input_path, data_only=True)
        ws_input = wb_input.active

        croquis_dir = os.path.join(temp_dir, 'croquis')
        os.makedirs(croquis_dir, exist_ok=True)

        generated_count = 0

        for row_idx in range(2, ws_input.max_row + 1):
            row_has_data = any(
                ws_input.cell(row=row_idx, column=col).value is not None
                for col in range(1, 9)
            )
            if not row_has_data:
                continue

            valores = {}
            for col_offset, celula_destino in CROQUI_MAPPINGS:
                col_origem = col_offset + 1
                value = ws_input.cell(row=row_idx, column=col_origem).value
                if value is not None:
                    valores[celula_destino] = value

            obra_value = ws_input.cell(row=row_idx, column=3).value
            if obra_value:
                filename = sanitize_filename(str(obra_value))
                if not filename.endswith('.xlsx'):
                    filename += '.xlsx'
                output_path = os.path.join(croquis_dir, filename)
                preencher_modelo_xlsx(MODELO_CROQUI, valores, output_path)
                generated_count += 1

        wb_input.close()

        if generated_count == 0:
            return jsonify({'error': 'Nenhum dado válido para processar'}), 400

        data_completa = datetime.now().strftime("%d-%m-%Y")

        zip_buffer = io.BytesIO()
        with ZipFile(zip_buffer, 'w') as zipf:
            for fname in os.listdir(croquis_dir):
                zipf.write(os.path.join(croquis_dir, fname), fname)
        zip_buffer.seek(0)

        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'croquis_gerados_{data_completa}.zip'
        )

    except Exception as e:
        return jsonify({'error': f'Erro ao processar arquivo: {str(e)}'}), 500

    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)